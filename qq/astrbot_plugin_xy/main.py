"""小月 QQ 接入插件 v1.4.1（借壳小白，第 8 课 + 第 9 课多人支持）。

路由规则（隐私优先）：
- 群聊：一律静默且 stop_event（v1.4.1 兜底——AstrBot 会话白名单关闭后
  群聊事件到达所有插件，这里先阻断后续插件响应）
- 私聊：任何 QQ 用户都能聊（v1.4 多人支持）——透传 sender QQ 号给
  小月服务 /api/chat，服务端按 QQ 号完全隔离记忆
- 陌生私聊：可聊，但仅限对话；主人专属功能（执行器/提醒/文件入库等）只在主人会话生效
- 文件入库：仅主人私聊可用（should_handle 白名单，与 v1.3 相同）

v1.4.1：
- AstrBot enable_id_white_list 必须为 False（否则陌生人私聊在
  whitelist_check 阶段被 event.stop_event 拦截，插件根本收不到）；
  群聊静默改由本插件 stop_event 兜底

v1.4 多人支持：
- /api/chat 请求带 user_id=sender（QQ 号），小月按人隔离记忆/事实/目标
- 群聊静默、owner 未配置 fail-closed 不变；访客不 stop_event（不挡其他插件）

v1.3 新增：主人发文件 → 自动入库知识库
- 支持 .txt/.md/.csv（直读）与 .docx/.pdf（解析提取）
- 大小上限 10MB；同名文件覆盖旧版（ingest 按 doc_name 幂等）
- 仅主人私聊可用（should_handle 白名单，与文本消息同闸门）
- 临时文件入库后即删

v1.1 修复（相对 v1.0）：
- 用 @filter.event_message_type(ALL) 消息处理器而不是 on_llm_request——
  AstrBot v4.27 的 on_llm_request 钩子跑在 LLM 门控之后，拦不住默认 LLM；
  消息处理器在 ProcessStage 门控之前执行，拦得住
- should_call_llm 语义：True=禁止默认 LLM（v1.0 用 False 是反的）
- 回复走 event.send()（置 _has_send_oper，双保险跳过默认 LLM）
"""
from __future__ import annotations

import asyncio
import base64
import hashlib
import hmac
import inspect
import os
import re
import tempfile
import time
import uuid
from pathlib import Path

import httpx
from astrbot.api import logger, AstrBotConfig
from astrbot.api.event import AstrMessageEvent, filter
from astrbot.api.message_components import Plain
from astrbot.api.star import Context, Star, register
from astrbot.core.message.message_event_result import MessageChain

REPLY_MAX_CHARS = 4000  # QQ 单条消息安全长度，超长截断并提示
FILE_MAX_BYTES = 10 * 1024 * 1024  # 文件入库上限 10MB（图文 PDF 常超 2MB，文本提取成本低）
FILE_TOTAL_TIMEOUT = 35  # 单个文件从取回到入库的总预算，避免 QQ 长时间无回复
VISION_MAX_IMAGE_BYTES = 10 * 1024 * 1024
VISION_TIMEOUT = 90
TEXT_EXTS = {".txt", ".md", ".markdown", ".csv", ".log", ".json"}
_IMAGE_MIME_BY_MAGIC = {
    "jpeg": "image/jpeg",
    "png": "image/png",
    "webp": "image/webp",
}
_ALLOWED_IMAGE_MIMES = frozenset(_IMAGE_MIME_BY_MAGIC.values())
_IMAGE_PLACEHOLDER_RE = re.compile(
    r"(?:\[\s*(?:image|img|图片|图像)(?:\s*:[^\]]*)?\s*\]|<\s*(?:image|img|图片|图像)(?:\s*:[^>]*)?>)",
    re.IGNORECASE,
)
_URL_RE = re.compile(r"https?://[^\s<>]+", re.IGNORECASE)


class ImageTooLargeError(Exception):
    """图片超过视觉请求大小上限。"""


class ImageFormatError(Exception):
    """图片不是允许的 JPEG/PNG/WebP，或 MIME 与文件头不一致。"""


class ImageDownloadError(Exception):
    """图片无法从 QQ 通道读取。"""


def clean_image_caption(value: str | None) -> str:
    """去掉 AstrBot 图片占位符和 URL，绝不把图片地址当作 caption。"""
    text = str(value or "")
    text = _IMAGE_PLACEHOLDER_RE.sub(" ", text)
    text = _URL_RE.sub(" ", text)
    return " ".join(text.split()).strip()


def detect_image_mime(prefix: bytes, content_type: str | None = None) -> str:
    """按文件头识别图片，并校验 HTTP MIME（缺失 MIME 时按本地文件头识别）。"""
    kind = ""
    if prefix.startswith(b"\xff\xd8\xff"):
        kind = "jpeg"
    elif prefix.startswith(b"\x89PNG\r\n\x1a\n"):
        kind = "png"
    elif len(prefix) >= 12 and prefix[:4] == b"RIFF" and prefix[8:12] == b"WEBP":
        kind = "webp"
    if not kind:
        raise ImageFormatError("仅支持 JPEG、PNG、WebP 图片")
    mime = _IMAGE_MIME_BY_MAGIC[kind]
    if content_type:
        actual = content_type.split(";", 1)[0].strip().lower()
        if actual not in _ALLOWED_IMAGE_MIMES or actual != mime:
            raise ImageFormatError("图片 MIME 与文件格式不匹配")
    return mime


def sign_qq_identity(secret: str, user_id: str, timestamp: str | int, request_id: str) -> str:
    """与服务端一致的 QQ 身份 HMAC-SHA256 签名。"""
    payload = f"{str(user_id).strip()}\n{str(timestamp).strip()}\n{str(request_id).strip()}".encode("utf-8")
    return hmac.new(str(secret).encode("utf-8"), payload, hashlib.sha256).hexdigest()


def build_qq_identity_headers(
    secret: str,
    user_id: str,
    *,
    request_id: str | None = None,
    timestamp: str | int | None = None,
) -> dict[str, str]:
    """生成带唯一 request_id 的 QQ 身份请求头。"""
    request_id = request_id or uuid.uuid4().hex
    timestamp = str(int(time.time()) if timestamp is None else timestamp)
    signature = sign_qq_identity(secret, user_id, timestamp, request_id)
    return {
        "X-QQ-User-ID": str(user_id).strip(),
        "X-QQ-Timestamp": timestamp,
        "X-QQ-Request-ID": request_id,
        # 同时发送通用头，方便网关/日志关联；服务端优先使用 QQ 专用头。
        "X-Request-ID": request_id,
        "X-QQ-Signature": signature,
    }


# 旧测试/外部插件脚本可能使用更短的函数名。
make_identity_headers = build_qq_identity_headers


class FileTooLargeError(Exception):
    """NapCat 已提供文件大小，尚未下载就超过入库上限。"""
# 文本提取失败时的安全文件名（防路径穿越/非法字符）
_SAFE_NAME = re.compile(r'[\\/:*?"<>|]')


def find_file_id(history_messages: list, name: str) -> tuple[str, int | None] | None:
    """按文件名精确找最近一条 file_id 和大小；找不到绝不猜其他文件。"""
    for m in reversed(history_messages):
        for seg in m.get("message") or []:
            if seg.get("type") != "file":
                continue
            data = seg.get("data") or {}
            if data.get("file") != name:
                continue
            fid = data.get("file_id")
            if not fid:
                return None
            try:
                size = int(data.get("file_size")) if data.get("file_size") else None
            except (TypeError, ValueError):
                size = None
            return str(fid), size
    return None


# NapCat 容器路径 → 宿主机路径映射的默认值（可被 container_path_map 配置覆盖）
_DEFAULT_PATH_MAP = (
    ("/app/.config/QQ/", "/opt/napcat/qq_config/"),
    ("/app/napcat/cache/", "/opt/napcat/cache/"),
    ("/app/napcat/config/", "/opt/napcat/config/"),
)


def _parse_path_map(cfg_value: str) -> tuple:
    """配置串 '容器前缀=宿主机前缀;...' → 映射元组；空/格式错回落默认。"""
    if not str(cfg_value or "").strip():
        return _DEFAULT_PATH_MAP
    pairs = []
    for seg in str(cfg_value).split(";"):
        if "=" in seg:
            cont, host = seg.split("=", 1)
            if cont.strip() and host.strip():
                pairs.append((cont.strip(), host.strip()))
    return tuple(pairs) if pairs else _DEFAULT_PATH_MAP


def to_host_path(p: str, path_map=_DEFAULT_PATH_MAP) -> str:
    """容器内路径翻译为宿主机路径（挂载点映射）；非容器路径原样返回。"""
    for cont, host in path_map:
        if p.startswith(cont):
            return host + p[len(cont):]
    return p


def should_handle(sender: str, group: str, owner_qq: str) -> bool:
    """主人专属闸门（纯函数，可单测）——文件入库等主人功能用。

    规则：群聊一律不处理（隐私铁律）；owner 未配置=全拒（fail-closed）；
    仅主人私聊返回 True。
    """
    if group:
        return False
    owner = str(owner_qq or "").strip()
    if not owner or str(sender or "").strip() != owner:
        return False
    return True


def can_chat(sender: str, group: str, owner_qq: str) -> bool:
    """聊天放行判定（v1.4 多人支持，纯函数，可单测）。

    规则：群聊一律不处理；owner 未配置=全拒（fail-closed，与 v1.3 一致）；
    私聊（含陌生人）一律放行——小月服务端按 sender QQ 号隔离记忆。
    """
    if group:
        return False
    if not str(owner_qq or "").strip():
        return False  # owner 未配置：无法区分主人/文件权限，宁可全拒
    return bool(str(sender or "").strip())


def extract_text(path: str) -> tuple[str, str]:
    """从文件提取纯文本。返回 (文本, 错误信息)，成功时错误信息为空。

    支持：txt/md/csv/log/json 直读；docx（python-docx）；pdf（pypdf）。
    """
    ext = Path(path).suffix.lower()
    try:
        if ext in TEXT_EXTS:
            text = Path(path).read_text(encoding="utf-8", errors="replace")
            return text, ""
        if ext == ".docx":
            from docx import Document  # python-docx，服务器同款依赖

            doc = Document(path)
            paras = [p.text for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:  # 表格文本也要（设定卡常见表格）
                for row in table.rows:
                    cells = [c.text.strip() for c in row.cells if c.text.strip()]
                    if cells:
                        paras.append(" | ".join(cells))
            return "\n".join(paras), ""
        if ext == ".pdf":
            from pypdf import PdfReader

            reader = PdfReader(path)
            pages = [(page.extract_text() or "") for page in reader.pages]
            return "\n".join(pages), ""
        return "", f"暂不支持 {ext or '无扩展名'} 格式（支持 txt/md/csv/docx/pdf）"
    except Exception as e:
        return "", f"解析失败：{type(e).__name__}: {e}"


def safe_doc_name(name: str) -> str:
    """文件名 → 安全文档名（去路径成分与非法字符，防穿越）。"""
    # QQ 端文件名可能携带 Windows 反斜杠路径（..\\..\\evil.txt），
    # 先统一成 / 再取 basename——Linux 上 Path().name 不认反斜杠分隔符
    name = name.replace("\\", "/")
    name = Path(name).name  # 去目录成分
    return _SAFE_NAME.sub("_", name).strip() or "未命名文档"


@register(
    "astrbot_plugin_xy",
    "小月接入",
    "小月 QQ 接入（借壳小白）：私聊直达小月服务（多人按 QQ 号隔离记忆），群聊静默",
    "v1.4.1",
)
class XiaoYuePlugin(Star):
    def __init__(self, context: Context, config: AstrBotConfig | None = None):
        super().__init__(context)
        self.cfg = config or {}
        # trust_env=False：本机回环调用不走系统代理——宿主机若有 HTTP_PROXY
        # 且 NO_PROXY 不含 127.0.0.1，Bearer token 会流经代理（全套已踩过的坑）
        self._client = httpx.AsyncClient(timeout=300, trust_env=False)  # 长文生成档 240s，留余量
        # QQ 文件 CDN 直连常 502（出网受限），下载兜底走代理（可配置，默认 clash）
        proxy = str(self.cfg.get("download_proxy", "") or "").strip()
        self._proxy_client = httpx.AsyncClient(
            timeout=120, trust_env=False, proxy=proxy or "http://127.0.0.1:7890"
        )
        try:
            self._vision_timeout_seconds = max(1.0, float(self.cfg.get("vision_timeout", VISION_TIMEOUT)))
        except (TypeError, ValueError):
            self._vision_timeout_seconds = float(VISION_TIMEOUT)
        try:
            self._vision_max_image_bytes = max(
                1, int(self.cfg.get("vision_max_image_bytes", VISION_MAX_IMAGE_BYTES))
            )
        except (TypeError, ValueError):
            self._vision_max_image_bytes = VISION_MAX_IMAGE_BYTES

    def _api_headers(self, sender: str, request_id: str | None = None) -> dict[str, str]:
        """构造服务端 Bearer + QQ 身份签名请求头。"""
        token = str(self.cfg.get("api_token", "") or "").strip()
        headers = {"Authorization": f"Bearer {token}"} if token else {}
        secret = str(self.cfg.get("identity_secret", "") or "").strip()
        if secret:
            headers.update(build_qq_identity_headers(secret, sender, request_id=request_id))
        return headers

    @filter.event_message_type(filter.EventMessageType.ALL)
    async def on_message(self, event: AstrMessageEvent):
        msg = event.get_message_str() or ""
        sender = event.get_sender_id() or ""
        group = event.get_group_id() or ""
        owner = str(self.cfg.get("owner_qq", "") or "").strip()

        # 群聊一律静默（隐私铁律）。v1.4.1：AstrBot 会话白名单已关闭
        # （否则陌生人私聊在 whitelist_check 阶段就被拦、到不了本插件），
        # 群聊事件因此会到达所有插件——这里 stop_event 兜底，阻止后续插件响应。
        if group:
            logger.debug("[xy] 群聊静默 stop_event group=%s sender=%s", group, sender)
            event.stop_event()
            event.should_call_llm(True)
            return
        # 私聊放行闸门（v1.4：陌生人可聊；owner 未配置 fail-closed）
        if not can_chat(sender, group, owner):
            # 仅禁止默认 LLM；不 stop_event，避免影响 meme_manager 等其他插件的事件处理。
            event.should_call_llm(True)
            return
        # 主人消息本插件全权处理，阻断其他处理器；访客不 stop_event（留给其他插件）
        if sender == owner:
            event.stop_event()

        # 文件分支（仅主人）：识别 File 组件 → 提取文本 → 入库知识库
        if should_handle(sender, group, owner):
            file_comp = self._find_file_component(event)
            if file_comp is not None:
                try:
                    await asyncio.wait_for(
                        self._handle_file(event, file_comp), timeout=FILE_TOTAL_TIMEOUT
                    )
                except asyncio.TimeoutError:
                    logger.warning("[xy] 文件处理超过总时限 %ss", FILE_TOTAL_TIMEOUT)
                    await event.send(MessageChain([Plain("❌ 文件处理超时（35秒），请稍后重发或先压缩文件")]))
                return

        # 图片只在 can_chat 门控之后处理；访客可问图，但不会进入主人专属分支。
        image_comp = self._find_image_component(event)
        if image_comp is not None:
            # 先阻断宿主默认 LLM，再做可能耗时的取图/上传；异常路径也不会漏处理。
            event.should_call_llm(True)
            await self._handle_image(event, image_comp, clean_image_caption(msg))
            return

        if not msg.strip():
            # 空白消息同样拦默认 LLM（v1.2：漏拦会漏进宿主默认 LLM）
            event.should_call_llm(True)
            return

        base = str(self.cfg.get("api_base", "") or "").strip().rstrip("/")
        token = str(self.cfg.get("api_token", "") or "").strip()
        if not base or not token:
            logger.error("[xy] 插件配置缺失（api_base/api_token 为空），请到 AstrBot 控制台配置")
            event.should_call_llm(True)
            return

        try:
            # v1.4：透传 sender QQ 号——小月按人隔离记忆；身份由签名头证明，
            # body.user_id 仅作服务端一致性校验。request_id 同时用于重试幂等。
            request_id = uuid.uuid4().hex
            r = await self._client.post(
                f"{base}/api/chat",
                json={"message": msg.strip(), "user_id": sender, "request_id": request_id},
                headers=self._api_headers(sender, request_id),
            )
            r.raise_for_status()
            reply = r.json().get("reply", "") or ""
        except Exception as e:
            logger.warning(f"[xy] 小月服务调用失败: {type(e).__name__}: {e}")
            reply = "😅 小月服务暂时不可达（服务器在重启？），稍后再试"

        reply = reply.strip()
        if len(reply) > REPLY_MAX_CHARS:
            reply = reply[:REPLY_MAX_CHARS] + "\n…（内容过长已截断，完整版去电脑面板看）"

        # 直接发送并禁止默认 LLM（_has_send_oper 双保险）
        await event.send(MessageChain([Plain(reply)]))
        event.should_call_llm(True)

    async def terminate(self):
        try:
            await self._client.aclose()
        except Exception:
            pass
        try:
            await self._proxy_client.aclose()
        except Exception:
            pass

    @staticmethod
    def _find_image_component(event: AstrMessageEvent):
        """遍历消息链找 AstrBot Image，兼容不同版本的组件类和 type 字段。"""
        try:
            from astrbot.api.message_components import Image
        except ImportError:
            Image = None
        for comp in event.get_messages() or []:
            if Image is not None:
                try:
                    if isinstance(comp, Image):
                        return comp
                except TypeError:
                    pass
            if str(getattr(comp, "type", "") or "").strip().lower() in {"image", "img", "图片"}:
                return comp
            class_name = comp.__class__.__name__.strip().lower()
            if class_name in {"image", "img"} or "image" in class_name:
                return comp
        return None

    @staticmethod
    def _image_value(comp, *names):
        """宽容读取 Image 组件字段，避免绑定单一 AstrBot 版本。"""
        for name in names:
            try:
                value = getattr(comp, name, None)
            except Exception:
                value = None
            if value not in (None, ""):
                return value
        return None

    def _image_limit_message(self) -> str:
        return f"图片超过上限 {self._vision_max_image_bytes / 1024 / 1024:.0f}MB，请压缩后再试"

    def _validate_image_path(self, path: str, content_type: str | None = None) -> str:
        p = Path(path)
        try:
            size = p.stat().st_size
        except OSError as exc:
            raise ImageDownloadError("图片文件不可读") from exc
        if size > self._vision_max_image_bytes:
            raise ImageTooLargeError(self._image_limit_message())
        try:
            with p.open("rb") as fh:
                prefix = fh.read(32)
        except OSError as exc:
            raise ImageDownloadError("图片文件不可读") from exc
        return detect_image_mime(prefix, content_type)

    def _write_image_bytes(self, value: bytes, dest: str) -> str:
        if len(value) > self._vision_max_image_bytes:
            raise ImageTooLargeError(self._image_limit_message())
        total = 0
        with open(dest, "wb") as fh:
            view = memoryview(value)
            for offset in range(0, len(view), 1024 * 1024):
                chunk = view[offset : offset + 1024 * 1024]
                total += len(chunk)
                if total > self._vision_max_image_bytes:
                    raise ImageTooLargeError(self._image_limit_message())
                fh.write(chunk)
        return self._validate_image_path(dest)

    async def _download_image(self, url: str, dest: str) -> str:
        """以流式方式下载图片；直连失败后复用已有代理通道。"""
        errors = []
        last_error = None
        for label, client in (("direct", self._client), ("proxy", self._proxy_client)):
            try:
                async with client.stream(
                    "GET", url, timeout=self._vision_timeout_seconds
                ) as response:
                    response.raise_for_status()
                    content_type = str(response.headers.get("content-type", "")).split(";", 1)[0].strip().lower()
                    if content_type not in _ALLOWED_IMAGE_MIMES:
                        raise ImageFormatError("服务端返回了不支持的图片 MIME")
                    raw_length = response.headers.get("content-length")
                    try:
                        if raw_length and int(raw_length) > self._vision_max_image_bytes:
                            raise ImageTooLargeError(self._image_limit_message())
                    except ValueError:
                        pass
                    total = 0
                    with open(dest, "wb") as fh:
                        async for chunk in response.aiter_bytes():
                            if not chunk:
                                continue
                            total += len(chunk)
                            if total > self._vision_max_image_bytes:
                                raise ImageTooLargeError(self._image_limit_message())
                            fh.write(chunk)
                return self._validate_image_path(dest, content_type)
            except ImageTooLargeError:
                Path(dest).unlink(missing_ok=True)
                raise
            except Exception as exc:
                last_error = exc
                errors.append(f"{label}: {type(exc).__name__}")
                Path(dest).unlink(missing_ok=True)
        if isinstance(last_error, ImageFormatError):
            raise last_error
        raise ImageDownloadError("图片下载失败（直连/代理均不可用）") from last_error

    @staticmethod
    def _source_candidates(value):
        """展开 get_file/NapCat 返回的路径、URL、bytes 或字典结果。"""
        if isinstance(value, dict):
            for key in ("file", "path", "file_path", "local_path", "url"):
                if value.get(key) not in (None, ""):
                    yield value[key]
            encoded = value.get("base64")
            if encoded:
                try:
                    yield base64.b64decode(encoded)
                except Exception:
                    return
            return
        if isinstance(value, (bytes, bytearray, memoryview, str, os.PathLike)):
            yield value
            return
        for key in ("file", "path", "file_path", "local_path", "url"):
            item = getattr(value, key, None)
            if item not in (None, ""):
                yield item

    async def _napcat_image_source(self, comp):
        """若 Image 仅提供 file_id，尝试 NapCat get_file 会话通道。"""
        file_id = self._image_value(comp, "file_id", "fileid")
        ob_url = str(self.cfg.get("onebot_http", "") or "").strip().rstrip("/")
        if not file_id or not ob_url:
            return None
        token = str(self.cfg.get("onebot_token", "") or "").strip()
        headers = {"Authorization": f"Bearer {token}"} if token else {}
        try:
            response = await self._client.post(
                f"{ob_url}/get_file",
                json={"file_id": str(file_id)},
                headers=headers,
                timeout=self._vision_timeout_seconds,
            )
            response.raise_for_status()
            data = response.json()
            if data.get("status") != "ok":
                return None
            return (data.get("data") or {})
        except Exception as exc:
            logger.warning("[xy] NapCat 图片 get_file 失败: %s: %s", type(exc).__name__, exc)
            return None

    async def _resolve_image_source(self, comp, dest: str):
        """按本地 file/path → url → get_file → NapCat 顺序取得本地文件。"""
        async def resolve_values(values):
            for value in values:
                for candidate in self._source_candidates(value):
                    if isinstance(candidate, (bytes, bytearray, memoryview)):
                        mime = self._write_image_bytes(bytes(candidate), dest)
                        return dest, mime, True
                    candidate = os.fspath(candidate) if isinstance(candidate, os.PathLike) else str(candidate)
                    if candidate.startswith(("http://", "https://")):
                        mime = await self._download_image(candidate, dest)
                        return dest, mime, True
                    if Path(candidate).is_file():
                        mime = self._validate_image_path(candidate)
                        return candidate, mime, False
            return None

        result = await resolve_values([
            self._image_value(comp, "file", "path", "file_path", "local_path")
        ])
        if result:
            return result
        result = await resolve_values([self._image_value(comp, "url", "image_url", "src")])
        if result:
            return result
        getter = getattr(comp, "get_file", None)
        if callable(getter):
            try:
                try:
                    value = getter(allow_return_url=True)
                except TypeError:
                    value = getter()
                if inspect.isawaitable(value):
                    value = await asyncio.wait_for(value, timeout=self._vision_timeout_seconds)
                result = await resolve_values([value])
                if result:
                    return result
            except Exception as exc:
                logger.warning("[xy] Image.get_file 失败: %s: %s", type(exc).__name__, exc)
        result = await resolve_values([await self._napcat_image_source(comp)])
        if result:
            return result
        raise ImageDownloadError("未能读取图片（NapCat 未提供可用文件）")

    @staticmethod
    def _vision_error_reply(exc: Exception) -> str:
        if isinstance(exc, ImageTooLargeError):
            return f"❌ {exc}"
        if isinstance(exc, ImageFormatError):
            return f"❌ {exc}"
        if isinstance(exc, asyncio.TimeoutError) or isinstance(exc, httpx.TimeoutException):
            return "❌ 图片处理超时（90秒），请稍后重试"
        if isinstance(exc, ImageDownloadError):
            return f"❌ {exc}"
        response = getattr(exc, "response", None)
        status = getattr(response, "status_code", None)
        if status == 400:
            return "❌ 图片文件无效或格式不支持，仅接受 JPEG、PNG、WebP"
        if status == 413:
            return "❌ 图片超过服务端大小上限，请压缩后再试"
        if status in {415, 422}:
            return "❌ 图片格式不支持，仅接受 JPEG、PNG、WebP"
        if status in {408, 504}:
            return "❌ 图片处理超时，请稍后重试"
        if status in {401, 403}:
            return "❌ 图片请求被服务端拒绝，请检查 QQ 身份配置"
        if status == 429:
            return "❌ 图片请求过于频繁，请稍后再试"
        return "❌ 图片服务暂时不可用，请稍后再试"

    async def _handle_image(self, event: AstrMessageEvent, comp, caption: str) -> None:
        """读取并校验 Image 后以本地 multipart 发送到视觉 API。"""
        tmp_path = ""
        try:
            suffix = Path(str(self._image_value(comp, "name", "file_name") or "image.bin")).suffix or ".bin"
            fd, tmp_path = tempfile.mkstemp(prefix="xy_vision_", suffix=suffix)
            os.close(fd)
            local, mime, created = await self._resolve_image_source(comp, tmp_path)
            if not created:
                # 本地源文件只读不删；占位临时文件仍需清理。
                Path(tmp_path).unlink(missing_ok=True)
                tmp_path = ""
            request_id = uuid.uuid4().hex
            sender = str(getattr(event, "get_sender_id", lambda: "")() or "").strip()
            base = str(self.cfg.get("api_base", "") or "").strip().rstrip("/")
            token = str(self.cfg.get("api_token", "") or "").strip()
            if not base or not token:
                raise ImageDownloadError("插件配置缺失（api_base/api_token）")
            with open(local, "rb") as image_file:
                response = await self._client.post(
                    f"{base}/api/chat/vision",
                    data={"message": caption, "request_id": request_id, "user_id": sender},
                    files={"image": (Path(local).name, image_file, mime)},
                    headers=self._api_headers(sender, request_id),
                    timeout=self._vision_timeout_seconds,
                )
            response.raise_for_status()
            reply = (response.json().get("reply", "") or "").strip()
            if not reply:
                reply = "（图片已收到，但视觉服务没有返回文字）"
            if len(reply) > REPLY_MAX_CHARS:
                reply = reply[:REPLY_MAX_CHARS] + "\\n…（内容过长已截断，完整版去电脑面板看）"
            await event.send(MessageChain([Plain(reply)]))
        except Exception as exc:
            logger.warning("[xy] 图片处理失败: %s: %s", type(exc).__name__, exc)
            await event.send(MessageChain([Plain(self._vision_error_reply(exc))]))
        finally:
            if tmp_path:
                try:
                    Path(tmp_path).unlink(missing_ok=True)
                except Exception:
                    pass
        event.should_call_llm(True)

    async def _download(self, url: str, dest: str) -> bool:
        """下载 QQ 文件 URL：直连优先，失败走 clash 代理兜底（20s 短超时，别把管线拖死）。"""
        for label, client in (("direct", self._client), ("proxy", self._proxy_client)):
            try:
                resp = await client.get(url, timeout=20)
                resp.raise_for_status()
                Path(dest).write_bytes(resp.content)
                return True
            except Exception as e:
                logger.warning(f"[xy] 文件下载失败({label}): {type(e).__name__}: {e}")
        return False

    async def _napcat_get_file(self, name: str) -> tuple[str, bool] | None:
        """NapCat 会话下载：历史 API 找 file_id → onebot get_file → 本地路径/base64。

        这是最可靠的通道——NapCat 用 QQ 会话取文件，不经公网 CDN（直连会 502/挂起）。
        返回本地文件路径；不可用返回 None。
        """
        ob_url = str(self.cfg.get("onebot_http", "") or "").strip().rstrip("/")
        ob_token = str(self.cfg.get("onebot_token", "") or "").strip()
        owner = str(self.cfg.get("owner_qq", "") or "").strip()
        if not ob_url or not owner:
            return None
        headers = {"Authorization": f"Bearer {ob_token}"} if ob_token else {}
        try:
            r = await self._client.post(
                f"{ob_url}/get_friend_msg_history",
                json={"user_id": int(owner), "count": 8},
                headers=headers,
                timeout=15,
            )
            r.raise_for_status()
            payload = r.json()
            if payload.get("status") != "ok":
                return None
            msgs = (payload.get("data") or {}).get("messages") or []
            file_info = find_file_id(msgs, name)
            if not file_info:
                return None
            fid, declared_size = file_info
            if declared_size is not None and declared_size > FILE_MAX_BYTES:
                raise FileTooLargeError(
                    f"{declared_size / 1024 / 1024:.1f}MB 超过上限 {FILE_MAX_BYTES / 1024 / 1024:.0f}MB"
                )
            r2 = await self._client.post(
                f"{ob_url}/get_file",
                json={"file_id": fid},
                headers=headers,
                timeout=90,
            )
            r2.raise_for_status()
            j = r2.json()
            if j.get("status") != "ok":
                return None
            d = j.get("data") or {}
            if d.get("file"):
                host_path = to_host_path(
                    str(d["file"]),
                    _parse_path_map(self.cfg.get("container_path_map", "")),
                )
                if Path(host_path).exists():
                    return host_path, False
            if d.get("base64"):
                import base64

                dest = str(Path(os.environ.get("TEMP", "/tmp")) / f"xy_ingest_{os.getpid()}_{name}")
                Path(dest).write_bytes(base64.b64decode(d["base64"]))
                return dest, True
            return None
        except FileTooLargeError:
            raise
        except Exception as e:
            logger.warning(f"[xy] NapCat get_file 失败: {type(e).__name__}: {e}")
            return None

    # ── 文件入库（仅主人）─────────────────────────────────

    @staticmethod
    def _find_file_component(event: AstrMessageEvent):
        """遍历消息组件链找 File 组件；无则 None。

        兼容两种形态：本地已落盘（file=路径）与 URL 下发（file=http(s) 链接）。
        """
        try:
            from astrbot.api.message_components import File
        except ImportError:
            logger.warning("[xy] 当前 AstrBot 版本无 File 组件，文件入库不可用")
            return None
        for comp in event.get_messages() or []:
            if isinstance(comp, File):
                return comp
        return None

    async def _handle_file(self, event: AstrMessageEvent, comp) -> None:
        """文件 → 提取文本 → /api/knowledge/ingest → 回执。任何失败都给主人明确原因。"""
        name = safe_doc_name(getattr(comp, "name", "") or "未命名文档")
        sender = str(getattr(event, "get_sender_id", lambda: "")() or "").strip()
        tmp_path = ""
        try:
            # ① 拿到本地文件：优先 NapCat 会话下载（get_file，QQ 会话通道最可靠），
            #    失败退回 URL 自下载（直连→clash 代理），再退回已落盘路径
            try:
                file_result = await self._napcat_get_file(name)
            except FileTooLargeError as e:
                await event.send(MessageChain([Plain(f"❌ 文件「{name}」{e}，请拆分后再发")]))
                return
            local = file_result[0] if file_result else None
            if file_result and file_result[1]:
                tmp_path = local
            if not local:
                src = ""
                try:
                    getter = getattr(comp, "get_file", None)
                    if callable(getter):
                        src = str(await getter(allow_return_url=True) or "")
                    if not src:
                        src = str(getattr(comp, "url", "") or "")
                except Exception as e:
                    logger.warning(f"[xy] get_file 失败: {type(e).__name__}: {e}")
                    src = ""
                if src.startswith(("http://", "https://")):
                    tmp_path = str(Path(os.environ.get("TEMP", "/tmp")) / f"xy_ingest_{os.getpid()}_{name}")
                    if not await self._download(src, tmp_path):
                        await event.send(MessageChain([Plain(
                            f"❌ 文件「{name}」下载失败（NapCat 会话与直连/代理均失败），稍后再发一次试试"
                        )]))
                        return
                    local = tmp_path
                elif src and Path(src).exists():
                    local = src
                else:
                    await event.send(MessageChain([Plain(f"❌ 文件「{name}」未能落盘（ NapCat 未提供路径），换个方式再发一次试试")]))
                    return

            # ② 大小限制
            size = Path(local).stat().st_size
            if size > FILE_MAX_BYTES:
                await event.send(MessageChain([Plain(f"❌ 文件「{name}」{size / 1024 / 1024:.1f}MB 超过上限 10MB，请拆分后再发")]))
                return

            # ③ 提取文本
            # PDF/docx 解析是同步 CPU/IO，放到线程避免阻塞 AstrBot 事件循环。
            text, err = await asyncio.to_thread(extract_text, local)
            if err:
                await event.send(MessageChain([Plain(f"❌ 「{name}」{err}")]))
                return
            if not text.strip():
                await event.send(MessageChain([Plain(f"❌ 「{name}」没有提取到文本（扫描版 PDF/空文件？）")]))
                return

            # ④ 入库（同名覆盖）
            base = str(self.cfg.get("api_base", "") or "").strip().rstrip("/")
            token = str(self.cfg.get("api_token", "") or "").strip()
            if not base or not token:
                await event.send(MessageChain([Plain("❌ 插件配置缺失（api_base/api_token），请到 AstrBot 控制台配置")]))
                return
            request_id = uuid.uuid4().hex
            r = await self._client.post(
                f"{base}/api/knowledge/ingest",
                json={"name": name, "content": text, "user_id": sender, "request_id": request_id},
                headers=self._api_headers(str(getattr(event, "get_sender_id", lambda: "")() or ""), request_id),
            )
            r.raise_for_status()
            chunks = r.json().get("chunks", 0)
            stem = Path(name).stem
            await event.send(MessageChain([Plain(
                f"📥 《{name}》已入库：{len(text)} 字，切了 {chunks} 块。\n"
                f"现在可以直接问我它的内容了（如「根据《{stem}》…」）"
            )]))
        except Exception as e:
            logger.warning(f"[xy] 文件入库失败 {name}: {type(e).__name__}: {e}")
            await event.send(MessageChain([Plain(f"❌ 「{name}」入库失败：{type(e).__name__}: {e}")]))
        finally:
            # 临时下载文件即删（源文件不动）
            if tmp_path:
                try:
                    Path(tmp_path).unlink(missing_ok=True)
                except Exception:
                    pass
