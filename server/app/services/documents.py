"""文档生成与保存：对话式"写文档"→ LLM 生成 → 存 documents 表 + 同步知识库。

保存后的文档自动 ingest 进知识库（切块+向量化），之后可检索/问答。
"""
import logging
import re
from datetime import datetime, timezone
from pathlib import Path

from openai import OpenAIError

from app.core import knowledge, llm
from app.models.database import connect

logger = logging.getLogger("assistant.documents")

DOC_PROMPT = """你是文档撰写助手。根据用户给出的标题与要求，生成一份结构化的 Markdown 文档。
要求：
- 标题用 # 开头
- 内容分小节（## 标题），条理清晰
- 基于用户要求写实，不编造用户没说过的信息
只输出文档正文（Markdown），不要多余解释。
"""


def parse_doc_command(msg: str) -> tuple[str, str] | None:
    """解析"写文档"命令 → (标题, 要求)。非命令返回 None。

    支持："写文档：标题XXX，内容：YYY" / "写文档 标题XXX" 等口语变体。
    """
    m = re.match(r"^(?:写文档|生成文档|写一份文档|帮我写文档)[:：]?\s*(.+)$", msg.strip())
    if not m:
        return None
    rest = m.group(1).strip()
    # "标题：X，内容：Y" / "标题X，内容：Y" / "X，内容：Y"
    title, requirement = rest, ""
    tm = re.match(r"标题[:：]?\s*(.+?)\s*[，,]\s*内容[:：]\s*(.+)$", rest)
    if tm:
        title, requirement = tm.group(1).strip(), tm.group(2).strip()
    else:
        cm = re.match(r"(.+?)\s*[，,]\s*内容[:：]\s*(.+)$", rest)
        if cm:
            title, requirement = cm.group(1).strip(), cm.group(2).strip()
    # 兜底：标题可能残留"标题"前缀（如"写文档：标题周报"无逗号时）
    title = re.sub(r"^标题[:：]?\s*", "", title).strip()
    return (title[:60], requirement or title)


async def generate_and_save(
    title: str,
    requirement: str,
    user_id: str | None = None,
    request_id: str | None = None,
) -> dict:
    """LLM 生成文档 → 存 documents 表 → 同步知识库。"""
    from app.core.memory import normalize_user_id
    from app.services.llm_usage import logical_request_id
    from app.services.sanitize import sanitize

    uid = normalize_user_id(user_id)
    logical_id = request_id or logical_request_id("document_generate", uid, "request")
    content = await llm.chat(
        [
            {"role": "system", "content": DOC_PROMPT},
            {"role": "user", "content": f"标题：{title}\n要求：{requirement}"},
        ],
        temperature=0.4,
        max_tokens=3000,
        timeout=240,  # 长文档生成档（全局 60s 会掐断长文）
        request_id=logical_id,
        user_id=uid,
    )
    content = sanitize(content.strip())
    if not content:
        return {"error": "生成失败：LLM 未返回内容"}
    title = sanitize(title)

    now = datetime.now(timezone.utc).isoformat()
    conn = connect()
    try:
        cur = conn.execute(
            "INSERT INTO documents (title, content, created_at) VALUES (?, ?, ?)",
            (title, content, now),
        )
        doc_id = cur.lastrowid
        conn.commit()
    finally:
        conn.close()

    # 同步进知识库（切块+向量化，之后可检索）
    try:
        await knowledge.ingest_document(f"文档-{title}", content)
    except (OpenAIError, TimeoutError, RuntimeError, ValueError, TypeError) as exc:
        logger.warning("文档知识库同步失败: %s", exc)

    return {"id": doc_id, "title": title, "words": len(content)}


def list_documents(limit: int = 20) -> list[dict]:
    conn = connect()
    try:
        rows = conn.execute(
            "SELECT id, title, created_at FROM documents ORDER BY id DESC LIMIT ?", (limit,)
        ).fetchall()
    finally:
        conn.close()
    return [dict(r) for r in rows]


def get_document(doc_id: int) -> dict | None:
    conn = connect()
    try:
        row = conn.execute(
            "SELECT id, title, content, created_at FROM documents WHERE id=?", (doc_id,)
        ).fetchone()
    finally:
        conn.close()
    return dict(row) if row else None


def export_docx(doc_id: int, out_dir: str = "") -> str:
    """把文档导出为 .docx（基础排版：标题居中加粗，## 转小标题）。

    返回文件绝对路径。依赖 python-docx。
    """
    from docx import Document as DocxDocument
    from docx.shared import Pt

    doc = get_document(doc_id)
    if not doc:
        raise FileNotFoundError(f"document {doc_id} not found")

    out_dir = Path(out_dir) if out_dir else Path(__file__).resolve().parents[2] / "data" / "exports"
    out_dir = out_dir.resolve()
    out_dir.mkdir(parents=True, exist_ok=True)
    safe_title = re.sub(r"[^\w\-.一-龥 ]", "_", Path(doc["title"]).name)[:20] or "document"
    out_path = (out_dir / f"doc-{doc_id}-{safe_title}.docx").resolve()
    if out_path.parent != out_dir:
        raise ValueError("导出路径超出导出目录")

    d = DocxDocument()
    title_p = d.add_paragraph()
    title_p.alignment = 1  # 居中
    run = title_p.add_run(doc["title"])
    run.bold = True
    run.font.size = Pt(16)

    for line in doc["content"].splitlines():
        line = line.strip()
        if not line:
            continue
        if line.startswith("## "):
            p = d.add_paragraph()
            r = p.add_run(line[3:])
            r.bold = True
            r.font.size = Pt(13)
        elif line.startswith("# "):
            continue  # 大标题已由 title 呈现
        elif line.startswith(("- ", "* ")):
            d.add_paragraph(line[2:], style="List Bullet")
        elif re.match(r"^\d+[.、]", line):
            d.add_paragraph(re.sub(r"^\d+[.、]\s*", "", line), style="List Number")
        else:
            d.add_paragraph(line)
    d.save(str(out_path))
    return str(out_path)
